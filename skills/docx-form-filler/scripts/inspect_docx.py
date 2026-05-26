#!/usr/bin/env python3
"""Dump docx structure: paragraphs (with index) + tables (with row/col cells).

Usage: python3 inspect_docx.py <file.docx>

Use this FIRST when filling a new template — it shows you which paragraph
indices contain underscores/placeholders and the cell layout of each table.
"""
import sys
from docx import Document

if len(sys.argv) != 2:
    print("Usage: inspect_docx.py <file.docx>", file=sys.stderr)
    sys.exit(1)

doc = Document(sys.argv[1])

print(f"=== {len(doc.paragraphs)} paragraphs ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip():
        print(f"P{i:3d}: {t!r}")

print(f"\n=== {len(doc.tables)} tables ===")
for ti, t in enumerate(doc.tables):
    print(f"\n--- Table {ti}: {len(t.rows)}r x {len(t.columns)}c ---")
    for ri, row in enumerate(t.rows):
        cells = [c.text.replace("\n", " | ")[:60] for c in row.cells]
        print(f"  R{ri}: {cells}")

# Headers / footers
for si, sec in enumerate(doc.sections):
    for hname in ("header", "first_page_header", "even_page_header",
                  "footer", "first_page_footer"):
        h = getattr(sec, hname, None)
        if h is None:
            continue
        for pi, p in enumerate(h.paragraphs):
            if p.text.strip():
                print(f"Sec{si}.{hname}.P{pi}: {p.text!r}")
