#!/usr/bin/env python3
"""TEMPLATE — copy this and adapt for each filling task.

Workflow:
  1. Run inspect_docx.py on the template — note paragraph indices + table layouts
  2. Run extract_pdf.py (or read source doc) — collect data
  3. Copy this file, edit SRC/DST and the data sections
  4. Run it
  5. Run inspect_docx.py on the OUTPUT to verify
  6. If textbox/cover content needs fixing, use xml_replace at the end
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from fill_helpers import set_para_text, set_cell_text, xml_replace

SRC = "/path/to/template.docx"
DST = "/path/to/output.docx"

doc = Document(SRC)
P = doc.paragraphs

# ---- Paragraph-based fields (personal info, simple lines) ----
# set_para_text(P[59], "Name：張三〔中文〕")
# set_para_text(P[61], "Zhang San〔English〕")

# ---- Table rows ----
# Each table column index is what inspect_docx.py reported.
# data = [(col1_value, col2_value, ...), ...]
# t = doc.tables[1]
# for i, (year, item, award) in enumerate(data):
#     row = t.rows[2 + i]  # skip header rows
#     set_cell_text(row.cells[1], year)
#     set_cell_text(row.cells[2], item)
#     set_cell_text(row.cells[3], award)
#     set_cell_text(row.cells[5], "")  # clear placeholder column

doc.save(DST)

# ---- Textbox / cover-page content (python-docx can't see these) ----
# xml_replace(DST, {
#     "PLACEHOLDER NAME": "Real Name",
# })

print(f"Saved: {DST}")
